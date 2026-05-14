"""
DX支援報告書のサンプルファイルを生成します（テスト用ダミーデータ）
  - DX報告書_001_製造業_在庫管理DX.docx
  - DX報告書_002_小売業_顧客分析DX.docx
  - DX報告書_003_物流業_配送最適化DX.pdf
  - DX報告書_004_医療機関_電子カルテDX.xlsx
"""
from pathlib import Path
import docx
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import mm

REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(exist_ok=True)

# ── Word ヘルパー ──────────────────────────────────────────
def add_title(doc, text):
    p = doc.add_heading(level=0)
    run = p.add_run(text)
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h1(doc, text):
    doc.add_heading(text, level=1)

def add_h2(doc, text):
    doc.add_heading(text, level=2)

def add_p(doc, text):
    doc.add_paragraph(text)

def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

# ── PDF ヘルパー ───────────────────────────────────────────
def build_pdf_styles(font_name):
    return {
        "title": ParagraphStyle("title", fontName=font_name, fontSize=18,
                                leading=26, spaceAfter=8, alignment=1),
        "h1":    ParagraphStyle("h1",    fontName=font_name, fontSize=14,
                                leading=20, spaceBefore=10, spaceAfter=4),
        "h2":    ParagraphStyle("h2",    fontName=font_name, fontSize=11,
                                leading=16, spaceBefore=6, spaceAfter=3),
        "body":  ParagraphStyle("body",  fontName=font_name, fontSize=9,
                                leading=14, spaceAfter=4),
        "bullet":ParagraphStyle("bullet",fontName=font_name, fontSize=9,
                                leading=14, leftIndent=12, spaceAfter=2),
    }

# ══════════════════════════════════════════════════════════
# 報告書 1: 製造業・在庫管理DX（Word）
# ══════════════════════════════════════════════════════════
def create_report_1():
    doc = docx.Document()
    add_title(doc, "DX支援報告書")

    add_h1(doc, "基本情報")
    add_p(doc, "報告書番号: DX-2024-001")
    add_p(doc, "支援期間: 2024年4月〜2024年9月（6ヶ月）")
    add_p(doc, "業種: 製造業（自動車向け金属部品製造）")
    add_p(doc, "企業規模: 従業員数250名、年商35億円")
    add_p(doc, "支援担当: 山田 太郎 / 鈴木 花子")

    add_h1(doc, "企業概要")
    add_p(doc, "株式会社〇〇精機は、自動車部品を中心とした金属部品の製造を行う中堅製造業です。愛知県に本社工場を置き、創業40年以上の製造実績を持ちます。主要取引先は国内大手自動車メーカー3社で、ISO9001取得済みの高い品質管理体制が強みです。一方、生産管理・在庫管理の面でデジタル化が遅れており、熟練担当者の経験と勘に依存した属人的な業務運営が課題となっていました。")

    add_h1(doc, "現状課題")

    add_h2(doc, "1. 在庫管理の非効率")
    add_p(doc, "在庫管理はExcelと紙台帳の併用で行われており、リアルタイムでの在庫把握が困難な状況でした。月次での棚卸しが必要で、欠品や過剰在庫が頻繁に発生していました。年間の在庫ロス額は約800万円に達しており、キャッシュフローへの影響も深刻な課題でした。")

    add_h2(doc, "2. 生産計画の属人化")
    add_p(doc, "生産計画は熟練担当者（在籍20年以上）の経験と勘に頼っており、担当者不在時に計画調整が困難でした。また、需要変動への対応が遅れ、急な増産・減産に対応できないケースが年に数回発生していました。2023年度は急な増産要請への対応遅れにより機会損失が約500万円発生しました。")

    add_h2(doc, "3. 品質データの一元管理不足")
    add_p(doc, "各工程（切削・研磨・熱処理・検査）の品質データが個別管理されており、不良品発生時のトレーサビリティが困難でした。品質問題の原因特定に平均3日以上かかっており、顧客からのクレーム対応も遅延が生じていました。")

    add_h1(doc, "DX施策・実施内容")

    add_h2(doc, "施策1: クラウドERP導入による在庫管理DX")
    add_p(doc, "製造業特化型クラウドERPシステム（SAP Business One）を導入しました。倉庫・工場の各拠点にIoTセンサーとバーコードリーダーを設置し、部品・製品の入出庫をリアルタイムで自動記録する仕組みを構築しました。スマートフォンによるQRコードスキャンで、現場担当者の負担を最小化しながら正確な在庫管理を実現しました。")

    add_h2(doc, "施策2: AIによる需要予測・生産計画最適化")
    add_p(doc, "過去5年分の受注データと生産実績をAIに学習させ、需要予測モデルを構築しました。機械学習アルゴリズム（XGBoost）を活用し、季節変動・取引先の生産計画・マクロ経済指標を考慮した高精度な予測を実現しました。このモデルを活用した生産計画支援ツールにより、担当者の経験に頼らない客観的な計画立案が可能になりました。")

    add_h2(doc, "施策3: 品質管理データの統合プラットフォーム構築")
    add_p(doc, "各工程の検査データをクラウドに集約する品質管理プラットフォームを構築しました。タブレット端末で現場から直接入力できる検査アプリを開発し、ペーパーレス化と同時にリアルタイムの品質監視を実現しました。問題発生時のトレーサビリティ機能により、ロット単位での原因特定が数時間以内に可能になりました。")

    add_h1(doc, "成果・効果")
    add_p(doc, "導入から6ヶ月後の測定結果は以下のとおりです。")
    add_bullet(doc, "在庫ロス額: 年間800万円 → 約250万円（68%削減）")
    add_bullet(doc, "在庫回転率: 年4.2回 → 年6.8回（62%向上）")
    add_bullet(doc, "生産計画立案時間: 週8時間 → 週2時間（75%削減）")
    add_bullet(doc, "品質問題の原因特定時間: 平均3日 → 平均4時間（87%短縮）")
    add_bullet(doc, "欠品発生件数: 月平均12件 → 月平均2件（83%削減）")
    add_bullet(doc, "システム投資回収期間: 約14ヶ月（試算）")

    add_h1(doc, "今後の展開")
    add_p(doc, "2025年度は取引先との受発注データ連携（EDI）を推進し、サプライチェーン全体の最適化を目指します。また、設備の予防保全IoTシステム（PdM: Predictive Maintenance）の導入も検討中です。製造業DXの第二フェーズとして、スマートファクトリー化に向けた取り組みを継続し、2026年度中のFAB全自動化ラインの試験稼働を目標としています。")

    path = REPORTS_DIR / "DX報告書_001_製造業_在庫管理DX.docx"
    doc.save(str(path))
    print(f"  生成完了: {path.name}")
    return path


# ══════════════════════════════════════════════════════════
# 報告書 2: 小売業・顧客分析DX（Word）
# ══════════════════════════════════════════════════════════
def create_report_2():
    doc = docx.Document()
    add_title(doc, "DX支援報告書")

    add_h1(doc, "基本情報")
    add_p(doc, "報告書番号: DX-2024-002")
    add_p(doc, "支援期間: 2024年7月〜2025年1月（7ヶ月）")
    add_p(doc, "業種: 小売業（ドラッグストアチェーン）")
    add_p(doc, "企業規模: 店舗数45店舗、従業員数680名、年商120億円")
    add_p(doc, "支援担当: 田中 一郎 / 佐藤 美咲")

    add_h1(doc, "企業概要")
    add_p(doc, "〇〇ドラッグ株式会社は、関東近郊に45店舗を展開するドラッグストアチェーンです。医薬品・日用品・化粧品を主力商品とし、地域密着型の運営で安定した顧客基盤を持っています。会員カード会員数は約18万人ですが、データ活用が進んでおらず、競合大手チェーンとの差別化が急務な状況でした。オムニチャネル化と顧客体験の向上が経営最重要課題として位置づけられていました。")

    add_h1(doc, "現状課題")

    add_h2(doc, "1. 顧客データの分散・未活用")
    add_p(doc, "会員カードシステムが老朽化しており、購買履歴データは蓄積されているものの分析活用ができていませんでした。各店舗のPOSデータと会員データが連携されておらず、顧客の購買行動を統合的に把握することが不可能な状態でした。会員データの更新頻度も低く、実態と乖離したデータが約30%存在していました。")

    add_h2(doc, "2. 販促施策の非効率")
    add_p(doc, "チラシ・DMによる販促が中心で、全会員に同じ内容を一律送付していました。販促費に対する費用対効果の測定が困難で、年間販促費用2億円のROIが不明確でした。セグメント別の購買傾向や季節性を考慮していない均一な販促により、顧客の関心度は低下傾向にありました。")

    add_h2(doc, "3. 在庫・発注管理の非効率")
    add_p(doc, "各店舗での発注は担当者の経験に依存しており、季節変動や地域特性を十分に考慮した発注ができていませんでした。医薬品の廃棄ロスと化粧品の欠品が同時に発生する状況が続いており、月次の廃棄ロスは全社合計1,200万円に達していました。")

    add_h1(doc, "DX施策・実施内容")

    add_h2(doc, "施策1: 顧客データ統合基盤（CDP）の構築")
    add_p(doc, "Customer Data Platform（Salesforce CDP）を導入し、POS・会員・ECサイト・アプリの購買データをリアルタイムで統合しました。データクレンジングと名寄せ処理により、18万人の会員データの精度を大幅に向上させました。顧客セグメンテーション機能を活用し、購買パターン・ライフステージ・健康意識などの軸で200以上のマイクロセグメントを定義しました。")

    add_h2(doc, "施策2: パーソナライズマーケティングの実現")
    add_p(doc, "LINE公式アカウントとCDPを連携し、顧客セグメントに応じた個別クーポン・商品情報・健康情報の配信を開始しました。AIによるレコメンデーションエンジンを実装し、個人の購買履歴・閲覧履歴・季節性・健康イベント（花粉シーズン等）を考慮した最適なタイミングでのパーソナライズ情報提供を実現しました。")

    add_h2(doc, "施策3: 需要予測による自動発注最適化")
    add_p(doc, "機械学習（LightGBM）による需要予測モデルを構築し、店舗・商品カテゴリ・SKU単位の発注推奨システムを導入しました。天候データ・近隣イベント・周辺人口動態・競合情報・季節要因を外部データとして取り込み、予測精度を向上させました。発注業務の自動化により、担当者は例外対応と品揃え企画に集中できる環境を整備しました。")

    add_h1(doc, "成果・効果")
    add_p(doc, "導入から6ヶ月後の測定結果は以下のとおりです。")
    add_bullet(doc, "会員のLINE友だち登録率: 8% → 34%（4.3倍）")
    add_bullet(doc, "パーソナライズDM開封率: 全体配信12% → 個別配信41%（3.4倍）")
    add_bullet(doc, "販促費ROI: 185% → 340%（84%改善）")
    add_bullet(doc, "廃棄ロス金額: 月平均1,200万円 → 月平均680万円（43%削減）")
    add_bullet(doc, "会員顧客単価: 月平均4,200円 → 月平均5,100円（21%向上）")
    add_bullet(doc, "リピート率（3ヶ月以内再来店）: 52% → 67%（15ポイント向上）")
    add_bullet(doc, "NPS（顧客推奨度）: +12 → +28（16ポイント向上）")

    add_h1(doc, "今後の展開")
    add_p(doc, "2025年度はスマートフォンアプリ開発によるオムニチャネル強化と、調剤薬局のオンライン処方箋サービス開始を予定しています。ヘルスケアデータ（健康診断結果・服薬履歴等）と購買データを連携した個人健康管理サポートサービスの構築も視野に入れ、単なる物販を超えた「地域の健康増進パートナー」としてのポジション確立を目指します。")

    path = REPORTS_DIR / "DX報告書_002_小売業_顧客分析DX.docx"
    doc.save(str(path))
    print(f"  生成完了: {path.name}")
    return path


# ══════════════════════════════════════════════════════════
# 報告書 3: 物流業・配送最適化DX（PDF）
# ══════════════════════════════════════════════════════════
def create_report_3():
    # Register Japanese font from Windows
    font_candidates = [
        ("C:/Windows/Fonts/msgothic.ttc", 0),
        ("C:/Windows/Fonts/YuGothR.ttc", 0),
        ("C:/Windows/Fonts/meiryo.ttc", 0),
        ("C:/Windows/Fonts/BIZ-UDGothicR.ttc", 0),
    ]
    font_name = None
    for font_path, idx in font_candidates:
        if Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont("JpFont", font_path, subfontIndex=idx))
                font_name = "JpFont"
                break
            except Exception:
                continue

    if font_name is None:
        print("  警告: 日本語フォントが見つかりません。PDF の代わりに Word ファイルを作成します。")
        return _create_report_3_docx_fallback()

    styles = build_pdf_styles(font_name)
    path = REPORTS_DIR / "DX報告書_003_物流業_配送最適化DX.pdf"

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    content = [
        Paragraph("DX支援報告書", styles["title"]),
        Spacer(1, 6 * mm),

        Paragraph("基本情報", styles["h1"]),
        Paragraph("報告書番号: DX-2024-003", styles["body"]),
        Paragraph("支援期間: 2024年10月〜2025年4月（7ヶ月）", styles["body"]),
        Paragraph("業種: 物流業（宅配・ラストマイル配送サービス）", styles["body"]),
        Paragraph("企業規模: 従業員数1,200名（ドライバー700名含む）、年商85億円", styles["body"]),
        Paragraph("支援担当: 中村 健一 / 渡辺 さくら", styles["body"]),
        Spacer(1, 4 * mm),

        Paragraph("企業概要", styles["h1"]),
        Paragraph(
            "〇〇デリバリー株式会社は、首都圏を中心に宅配・ラストマイル配送を行う物流企業です。"
            "EC市場の拡大に伴い取扱物量は5年で2.3倍に増加しましたが、ドライバー不足と配送コストの上昇が"
            "経営を圧迫していました。1日あたりの配送件数は約85,000件で、再配達率が22%と業界平均を上回る"
            "状態が続いており、効率化による持続可能な事業運営が急務でした。",
            styles["body"]
        ),
        Spacer(1, 4 * mm),

        Paragraph("現状課題", styles["h1"]),

        Paragraph("1. 配送ルートの非最適化", styles["h2"]),
        Paragraph(
            "配送ルートの計画は各ドライバーの経験と判断に依存しており、標準化されていませんでした。"
            "経験の浅いドライバーと熟練ドライバーで同エリアの配送効率に最大40%の差があり、"
            "走行距離の無駄が燃料コストを押し上げていました。渋滞情報・天候・荷物量の変動に"
            "リアルタイムで対応したルート更新ができていませんでした。",
            styles["body"]
        ),

        Paragraph("2. 再配達率の高止まり", styles["h2"]),
        Paragraph(
            "再配達率22%は業界平均（約15%）を大幅に上回り、運営コストの深刻な要因となっていました。"
            "配達時間帯の事前通知システムがなく、受取人が不在のまま配達するケースが多発していました。"
            "再配達1件あたりのコストは約450円で、年間の再配達コストは約8億円に達していました。",
            styles["body"]
        ),

        Paragraph("3. 車両・積載管理の非効率", styles["h2"]),
        Paragraph(
            "車両の積載率は平均62%にとどまり、空きスペースを抱えたまま走行するケースが多くありました。"
            "リアルタイムでの車両位置・積載状況の把握ができず、急な追加集荷依頼への対応も困難でした。"
            "車両の稼働状況・燃費・メンテナンス状況の一元管理も行われていませんでした。",
            styles["body"]
        ),
        Spacer(1, 4 * mm),

        Paragraph("DX施策・実施内容", styles["h1"]),

        Paragraph("施策1: AIによる動的配送ルート最適化システム", styles["h2"]),
        Paragraph(
            "機械学習と遺伝的アルゴリズムを組み合わせた配送ルート最適化エンジンを開発しました。"
            "リアルタイム交通情報（Google Maps Platform）・天候データ・荷物の優先度・"
            "配達時間帯指定情報を統合的に処理し、1ドライバーあたり数百件の配送先を"
            "最適な順序とルートで自動計画します。"
            "ルート変更が必要な場合も、走行中のドライバーのスマートフォンアプリへ"
            "リアルタイムに通知・更新します。",
            styles["body"]
        ),

        Paragraph("施策2: 配達時間帯事前通知・受取確認システム", styles["h2"]),
        Paragraph(
            "配達2時間前にSMS・LINEで受取人へ配達予定時刻を通知し、"
            "不在の場合は時間変更・置き配・宅配ボックス指定・コンビニ受取への"
            "切り替えをワンクリックで可能にするシステムを構築しました。"
            "AIによる不在予測モデルを活用し、不在リスクの高い配達から優先的に通知を送る"
            "スマート通知機能も実装しました。",
            styles["body"]
        ),

        Paragraph("施策3: IoTによる車両・積載リアルタイム管理", styles["h2"]),
        Paragraph(
            "全車両にIoT車載端末を設置し、位置情報・走行速度・急加速急ブレーキ・"
            "燃費・エンジン稼働状況をリアルタイムで収集するフリート管理システムを導入しました。"
            "荷室への重量センサー設置により積載率を自動計測し、空き容量がある車両への"
            "追加集荷依頼のマッチングも自動化しました。"
            "収集データを活用した予防保全アラートにより、故障による配送遅延を未然に防ぎます。",
            styles["body"]
        ),
        Spacer(1, 4 * mm),

        Paragraph("成果・効果", styles["h1"]),
        Paragraph("導入から7ヶ月後の測定結果は以下のとおりです。", styles["body"]),
        Paragraph("・ 1台あたり1日の配送件数: 平均42件 → 平均54件（29%向上）", styles["bullet"]),
        Paragraph("・ 再配達率: 22% → 11%（50%削減）", styles["bullet"]),
        Paragraph("・ 走行距離（1配送あたり）: 平均3.8km → 平均2.9km（24%削減）", styles["bullet"]),
        Paragraph("・ 配送コスト: 前年比22%削減（年間約18億円 → 約14億円）", styles["bullet"]),
        Paragraph("・ 燃料費: 前年比18%削減", styles["bullet"]),
        Paragraph("・ ドライバー時間外労働: 月平均38時間 → 月平均25時間（34%削減）", styles["bullet"]),
        Paragraph("・ 車両積載率: 62% → 79%（17ポイント向上）", styles["bullet"]),
        Paragraph("・ 顧客満足度（CS）スコア: 68点 → 84点（16点向上）", styles["bullet"]),
        Spacer(1, 4 * mm),

        Paragraph("今後の展開", styles["h1"]),
        Paragraph(
            "2025年度後半には、自動配送ロボット（UGV）の試験導入を都内2エリアで開始予定です。"
            "また、宅配ボックスのIoT化によるスマートロッカー連携、"
            "ドローン配送（一部山間エリア）の実証実験も計画しています。"
            "2027年度には完全自動化配送センターの稼働を目標とし、"
            "人手不足が深刻化する物流業界における持続可能なビジネスモデルの確立を目指します。",
            styles["body"]
        ),
    ]

    doc.build(content)
    print(f"  生成完了: {path.name}")
    return path


def _create_report_3_docx_fallback():
    """日本語フォントが利用できない場合のWord版フォールバック"""
    doc = docx.Document()
    add_title(doc, "DX支援報告書")
    add_h1(doc, "基本情報")
    add_p(doc, "報告書番号: DX-2024-003")
    add_p(doc, "支援期間: 2024年10月〜2025年4月（7ヶ月）")
    add_p(doc, "業種: 物流業（宅配・ラストマイル配送サービス）")
    add_h1(doc, "企業概要")
    add_p(doc, "〇〇デリバリー株式会社は、首都圏を中心に宅配・ラストマイル配送を行う物流企業です。EC市場の拡大に伴い取扱物量は5年で2.3倍に増加しましたが、ドライバー不足と配送コストの上昇が経営を圧迫していました。")
    add_h1(doc, "現状課題")
    add_p(doc, "配送ルートの非最適化: 各ドライバーの経験に依存したルート計画。再配達率22%と業界平均を上回る高止まり。車両積載率62%と低水準。")
    add_h1(doc, "DX施策")
    add_p(doc, "AIによる動的配送ルート最適化、配達時間帯事前通知システム、IoT車両管理システムを導入しました。")
    add_h1(doc, "成果・効果")
    add_bullet(doc, "配送件数: 42件/日 → 54件/日（29%向上）")
    add_bullet(doc, "再配達率: 22% → 11%（50%削減）")
    add_bullet(doc, "配送コスト: 前年比22%削減")
    add_bullet(doc, "ドライバー残業: 月38h → 月25h（34%削減）")
    path = REPORTS_DIR / "DX報告書_003_物流業_配送最適化DX.docx"
    doc.save(str(path))
    print(f"  生成完了（Word版）: {path.name}")
    return path


# ══════════════════════════════════════════════════════════
# 報告書 4: 医療機関・電子カルテDX（Excel）
# ══════════════════════════════════════════════════════════
def create_report_4():
    wb = openpyxl.Workbook()

    header_font  = Font(bold=True, size=11)
    header_fill  = PatternFill("solid", fgColor="4472C4")
    header_font_white = Font(bold=True, size=11, color="FFFFFF")
    title_font   = Font(bold=True, size=14)
    section_fill = PatternFill("solid", fgColor="D9E1F2")

    def set_header(ws, row, col, value):
        cell = ws.cell(row=row, column=col, value=value)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        return cell

    def set_section(ws, row, col, value):
        cell = ws.cell(row=row, column=col, value=value)
        cell.font = Font(bold=True, size=10)
        cell.fill = section_fill
        return cell

    # ── シート1: 基本情報 ──
    ws1 = wb.active
    ws1.title = "基本情報"
    ws1.column_dimensions["A"].width = 22
    ws1.column_dimensions["B"].width = 45

    ws1.cell(row=1, column=1, value="DX支援報告書").font = title_font
    ws1.merge_cells("A1:B1")

    headers = [("項目", "内容")]
    data = [
        ("報告書番号",   "DX-2024-004"),
        ("支援期間",     "2024年5月〜2025年2月（10ヶ月）"),
        ("業種",         "医療機関（地域中核病院）"),
        ("企業規模",     "病床数320床、職員数850名（医師120名、看護師420名含む）"),
        ("年間患者数",   "外来患者数 約18万人、入院患者数 約6,800人"),
        ("支援担当",     "木村 誠司 / 高橋 理恵"),
        ("報告書作成日", "2025年3月15日"),
    ]
    set_header(ws1, 3, 1, "項目")
    set_header(ws1, 3, 2, "内容")
    for i, (k, v) in enumerate(data, start=4):
        ws1.cell(row=i, column=1, value=k).font = Font(bold=True)
        ws1.cell(row=i, column=2, value=v)

    # ── シート2: 現状課題 ──
    ws2 = wb.create_sheet("現状課題")
    ws2.column_dimensions["A"].width = 28
    ws2.column_dimensions["B"].width = 60

    ws2.cell(row=1, column=1, value="現状課題一覧").font = title_font
    set_header(ws2, 3, 1, "課題カテゴリ")
    set_header(ws2, 3, 2, "詳細")

    issues = [
        ("1. 紙カルテ・紙伝票の大量運用",
         "外来・入院ともに紙カルテが中心で、院内搬送に1件あたり平均15分を要していた。"
         "カルテ紛失リスクがあり、年間50件以上の紛失インシデントが発生。"
         "保管スペースとして別棟倉庫を賃借しており、年間コスト約600万円。"),
        ("2. 部門間情報共有の遅延",
         "検査・放射線・薬剤・看護の各部門がシステム非連携のため、"
         "検査結果の医師への通知に平均2時間以上かかっていた。"
         "緊急検査値（パニック値）の連絡も電話のみで、見落としリスクが存在。"),
        ("3. 医師の事務作業負担",
         "診療録記載・処方入力・紹介状作成などの事務作業に医師の勤務時間の約35%を占めていた。"
         "時間外勤務の主因となっており、医師の離職率が地域平均より8ポイント高い状態。"),
        ("4. 薬剤処方ミス・重複投与リスク",
         "複数科受診患者の処方情報が一元化されておらず、薬剤師によるチェックが属人的だった。"
         "年間の潜在的薬剤インシデント件数は院内調査で推計130件以上。"),
        ("5. 経営データの可視化不足",
         "診療科別・疾患別の収益性、病床稼働率、在院日数の分析が月次でしか行えなかった。"
         "経営判断のタイムラグが大きく、補助金・診療報酬の取り漏れが年間1,200万円と推計。"),
    ]
    for i, (cat, detail) in enumerate(issues, start=4):
        ws2.cell(row=i, column=1, value=cat).font = Font(bold=True)
        ws2.cell(row=i, column=2, value=detail).alignment = Alignment(wrap_text=True)
        ws2.row_dimensions[i].height = 60

    # ── シート3: DX施策 ──
    ws3 = wb.create_sheet("DX施策")
    ws3.column_dimensions["A"].width = 30
    ws3.column_dimensions["B"].width = 20
    ws3.column_dimensions["C"].width = 55

    ws3.cell(row=1, column=1, value="DX施策・実施内容").font = title_font
    set_header(ws3, 3, 1, "施策名")
    set_header(ws3, 3, 2, "導入ツール/技術")
    set_header(ws3, 3, 3, "実施内容")

    measures = [
        ("施策1: 電子カルテシステム全面導入",
         "HOPE LifeMark-HIS",
         "全診療科の電子カルテ化を実施。タブレット端末を病棟・外来に400台配備し、"
         "ベッドサイドでのリアルタイム入力を実現。音声入力AI（Nuance DAX）連携により"
         "医師の記録時間を大幅に短縮。文書テンプレートの整備でサマリ作成を自動化。"),
        ("施策2: 部門間情報統合基盤の構築",
         "HL7 FHIR / API連携",
         "検査・放射線・薬剤・栄養・リハビリ各部門システムをFHIR標準でAPI連携。"
         "検査結果の自動プッシュ通知をスマートフォンアプリで実現。"
         "パニック値は最優先アラートとして担当医・当直医に同時通知。"),
        ("施策3: AI服薬支援・処方チェック",
         "PHILOS（AI処方監査）",
         "AIによる処方内容の自動監査システムを導入。"
         "患者の全処方歴・アレルギー・腎機能データを参照し、"
         "禁忌・相互作用・重複投与をリアルタイムで検出・警告。"
         "薬剤師の最終確認フローと組み合わせてダブルチェック体制を確立。"),
        ("施策4: 経営ダッシュボード・BI導入",
         "Tableau / Power BI",
         "診療科別収益・病床稼働・DPC分析・在院日数・紹介逆紹介率を"
         "リアルタイムで可視化するダッシュボードを構築。"
         "経営会議での活用により、月次から週次への意思決定サイクル短縮。"
         "診療報酬の算定漏れをAIで自動検出するシステムも同時導入。"),
    ]
    for i, (施策, tool, detail) in enumerate(measures, start=4):
        ws3.cell(row=i, column=1, value=施策).font = Font(bold=True)
        ws3.cell(row=i, column=2, value=tool)
        ws3.cell(row=i, column=3, value=detail).alignment = Alignment(wrap_text=True)
        ws3.row_dimensions[i].height = 75

    # ── シート4: 成果・効果 ──
    ws4 = wb.create_sheet("成果・効果")
    ws4.column_dimensions["A"].width = 32
    ws4.column_dimensions["B"].width = 20
    ws4.column_dimensions["C"].width = 20
    ws4.column_dimensions["D"].width = 18

    ws4.cell(row=1, column=1, value="導入効果（10ヶ月後測定）").font = title_font
    set_header(ws4, 3, 1, "指標")
    set_header(ws4, 3, 2, "導入前")
    set_header(ws4, 3, 3, "導入後")
    set_header(ws4, 3, 4, "改善率")

    results = [
        ("カルテ記録時間（医師/日）",     "平均95分",     "平均38分",     "60%削減"),
        ("検査結果通知時間",               "平均2.2時間",  "平均8分",      "94%短縮"),
        ("処方インシデント発生件数/年",    "推計130件以上","推計28件",     "78%削減"),
        ("紙カルテ関連コスト/年",          "約1,400万円",  "約150万円",    "89%削減"),
        ("病床稼働率",                     "74.2%",        "82.7%",        "8.5pt向上"),
        ("平均在院日数",                   "14.8日",       "12.3日",       "17%短縮"),
        ("診療報酬算定漏れ回収額/月",      "不明",         "約280万円",    "新規回収"),
        ("医師の時間外労働時間/月平均",    "68時間",       "44時間",       "35%削減"),
        ("職員満足度スコア（5点満点）",    "2.8点",        "3.9点",        "1.1pt向上"),
        ("システム投資回収期間",           "―",            "推計22ヶ月",   "―"),
    ]
    for i, (metric, before, after, rate) in enumerate(results, start=4):
        ws4.cell(row=i, column=1, value=metric)
        ws4.cell(row=i, column=2, value=before).alignment = Alignment(horizontal="center")
        ws4.cell(row=i, column=3, value=after).alignment  = Alignment(horizontal="center")
        ws4.cell(row=i, column=4, value=rate).alignment   = Alignment(horizontal="center")

    # ── シート5: 今後の展開 ──
    ws5 = wb.create_sheet("今後の展開")
    ws5.column_dimensions["A"].width = 20
    ws5.column_dimensions["B"].width = 65

    ws5.cell(row=1, column=1, value="今後の展開ロードマップ").font = title_font
    set_header(ws5, 3, 1, "時期")
    set_header(ws5, 3, 2, "計画内容")

    roadmap = [
        ("2025年度上半期",
         "患者向けスマートフォンアプリ開発：予約・問診票・検査結果閲覧・会計のオンライン化。"
         "PHR（個人健康記録）との連携による継続的な健康管理支援サービスの開始。"),
        ("2025年度下半期",
         "AIによる疾患予測・重症化リスクスコアリングの試験導入。"
         "退院支援・地域連携パスのデジタル化による他医療機関・介護施設との情報共有基盤構築。"),
        ("2026年度",
         "手術室・ICUへのIoTセンサー拡充による患者モニタリング高度化。"
         "ロボティクス（自動搬送ロボット・自動調剤ロボット）の本格導入。"
         "地域医療情報ネットワーク（HPKI）への参加と地域全体のデータ連携推進。"),
        ("2027年度以降",
         "生成AIを活用した診療支援（診断補助・治療方針提案）の段階的実装。"
         "デジタルツインによる病院経営シミュレーション基盤の整備。"
         "スマートホスピタル認定取得を目標とした全院DX完成フェーズ。"),
    ]
    for i, (timing, plan) in enumerate(roadmap, start=4):
        ws5.cell(row=i, column=1, value=timing).font = Font(bold=True)
        ws5.cell(row=i, column=2, value=plan).alignment = Alignment(wrap_text=True)
        ws5.row_dimensions[i].height = 65

    path = REPORTS_DIR / "DX報告書_004_医療機関_電子カルテDX.xlsx"
    wb.save(str(path))
    print(f"  生成完了: {path.name}")
    return path


# ══════════════════════════════════════════════════════════
# メイン
# ══════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("=== サンプル報告書を生成中 ===\n")
    create_report_1()
    create_report_2()
    create_report_3()
    create_report_4()
    print(f"\n完了: {REPORTS_DIR}/ に4件のサンプルファイルを生成しました")
    print("次のステップ: python ingest.py")
